import os
import openai
import faiss
import json
import numpy as np
import streamlit as st
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
import base64
import fitz  # PyMuPDF for extracting text from PDFs
from docx import Document  # for handling Word documents

# Load environment variables
load_dotenv()

st.set_page_config(page_title="Regulaider", page_icon="🏥", layout="wide")

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Load FAISS index from file
index = faiss.read_index("faiss_index.index")

# Load document metadata with embeddings
with open('document_metadata_with_embeddings.json', 'r') as f:
    documents = json.load(f)

# Initialize HuggingFace embedding model
embedding_model = HuggingFaceEmbeddings()

# Function to add the logo in the sidebar
def add_sidebar_logo(logo_file):
    with open(logo_file, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode()

    st.sidebar.markdown(
        f"""
        <style>
        /* Style to position the logo in the sidebar */
        .sidebar-logo-container {{
            position: relative;
            width: 100%;
            padding-bottom: 20px;
        }}
        .sidebar-logo-container img {{
            width: 150px;
            height: auto;
        }}
        </style>
        <div class="sidebar-logo-container">
            <img src="data:image/png;base64,{encoded_image}" alt="Sidebar Logo">
        </div>
        """,
        unsafe_allow_html=True
    )

# Call the function to add the sidebar logo 
add_sidebar_logo("images/sidebar_logo.png")

# Function to add the logo in the top right corner
def add_logo(logo_file):
    with open(logo_file, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode()

    st.markdown(
        f"""
        <style>
        /* Position the logo in the top right corner */
        .logo-container {{
            position: absolute;
            top: 10px;
            right: -100px; 
            z-index: 1000;
        }}
        .logo-container img {{
            width: 150px;  
            height: auto;
        }}
        </style>
        <div class="logo-container">
            <img src="data:image/png;base64,{encoded_image}" alt="Logo">
        </div>
        """,
        unsafe_allow_html=True
    )

# Call the function with your image path
add_logo("images/logo.png")

# Helper function to generate a question from a chunk
def generate_question_from_chunk(chunk_text):
    prompt = f"Based on the following regulatory section, generate a critical audit-related question. Address the user personally (use 'you' or 'your'):\n\n{chunk_text}."

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a bot helping with audit preparation."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150,
        n=1
    )
    return response['choices'][0]['message']['content'].strip()

# Helper function to evaluate the user's response
def evaluate_response(user_response, chunk_text):
    # Step 1: Check if the user's response is relevant to the regulatory text
    relevance_check_prompt = (
        f"Determine if the following response is relevant to the given regulatory section. If the user's answers shows non attempt with phrases like I dont know or anything related then mark it as relevant\n\n"
        f"Regulatory Section:\n{chunk_text}\n\n"
        f"User Response:\n{user_response}\n\n"
        f"Respond with 'Relevant' if the response is on-topic and addresses the regulatory section, or "
        f"'Irrelevant' if it does not."
    )

    relevance_response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an assistant helping with audit preparation."},
            {"role": "user", "content": relevance_check_prompt}
        ],
        max_tokens=10,
        n=1
    )
    relevance = relevance_response['choices'][0]['message']['content'].strip()

    # Step 2: Provide feedback based on relevance
    if relevance.lower() == "relevant":
        # Detailed evaluation if response is relevant
        prompt = (
            f"Regulatory Section:\n{chunk_text}\n\n"
            f"Your Response:\n{user_response}\n\n"
            f"Feedback:\n"
            f"1. Evaluate whether your response complies with the regulatory section. Focus on key points where your response aligns or does not align with the requirements.\n"
            f"2. If your response does not fully meet the regulatory requirements, provide specific recommendations on what actions you need to take to achieve compliance. "
            f"Include detailed suggestions or steps you can follow to mitigate any issues."
        )

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an assistant helping with audit preparation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            n=1
        )

        return response['choices'][0]['message']['content'].strip()

    else:
        # Feedback if the response is irrelevant
        return "It seems that your response does not address the regulatory section directly. Please review the content carefully and ensure your response is relevant to the regulatory requirements specified."

# Helper function to generate a final evaluation after the conversation
def generate_final_evaluation(conversation):
    prompt = (
        "Address the user personally (use 'you' or 'your company' throughout). "
        "Based on the following conversation, evaluate your company's audit readiness. "
        "Use official FDA terms and summarize the observations as if you're an FDA investigator. "
        "Structure your response into the following sections:\n\n"
        
        "**Most Likely Outcome of the Audit:**\n"
        "Choose from one of the following outcomes: No Action Indicated (NAI), Voluntary Action Indicated (VAI), "
        "Form 483 (Inspectional Observations), or Warning Letter. "
        "Explain why this outcome applies to your company.\n\n"

        "**Key Observations:**\n"
        "Summarize key findings that support the chosen outcome. This should include both compliance and non-compliance issues observed during your audit preparation.\n\n"

        "**Recommendations:**\n"
        "Provide recommendations based on the outcome. If the outcome is positive (NAI or VAI), suggest how you and your company can improve further. "
        "If the outcome is negative (Warning Letter, Form 483), recommend how your company can mitigate the issues. "
        "If your company is already in compliance, give minimal suggestions for maintaining this status."
    )

    # Add the conversation to the prompt
    for message in conversation:
        prompt += f"{message['role'].capitalize()}: {message['content']}\n\n"

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a bot helping with audit preparation."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        n=1
    )

    return response['choices'][0]['message']['content'].strip()

# Helper function to retrieve a chunk and generate a question from FAISS index
def retrieve_chunk_and_generate_question(project_area_embedding):
    k = 5  # Top 5 most relevant chunks
    D, I = index.search(np.array([project_area_embedding]), k)

    # Filter out chunks that were already used
    unused_indices = [i for i in I[0] if i not in st.session_state.used_chunks]

    if not unused_indices:
        st.write("All relevant chunks have been used. Generating more questions may require new content.")
    else:
        relevant_chunk = documents[unused_indices[0]]
        st.session_state.relevant_chunk = relevant_chunk
        question = generate_question_from_chunk(relevant_chunk['page_content'])

        # Mark this chunk as used
        st.session_state.used_chunks.append(unused_indices[0])

        return question

# Function to save the conversation as a Word document
from datetime import datetime

def save_conversation_to_word(conversation, project_scope, filename_prefix="report"):
    doc = DocxDocument()
    
    # Add project scope as the heading
    doc.add_heading(f"{project_scope}", 0)

    # Add timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Created on: {timestamp}")
    
    for message in conversation:
        if message['role'] == 'bot':
            doc.add_paragraph("Regulaider:", style='Heading 2').bold = True
            doc.add_paragraph(message['content'])
        else:
            doc.add_paragraph("You:", style='Heading 2').bold = True
            doc.add_paragraph(message['content'])
        doc.add_paragraph("_" * 50)

    formatted_filename = f"{project_scope.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Save the document
    doc.save(formatted_filename)
    return formatted_filename

def save_conversation_to_pdf(conversation, project_scope, filename_prefix="report"):
    formatted_filename = f"{project_scope.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    pdf = SimpleDocTemplate(formatted_filename, pagesize=A4)
    elements = []

    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    bold_style = ParagraphStyle(name='Bold', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12)

    elements.append(Paragraph(f"{project_scope}", styles["Title"]))
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    elements.append(Paragraph(f"Created on: {timestamp}", normal_style))
    elements.append(Spacer(1, 0.2 * inch))

    for message in conversation:
        if message['role'] == 'bot':
            elements.append(Paragraph("<b>Regulaider:</b>", bold_style))
            elements.append(Paragraph(message['content'], normal_style))
        else:
            elements.append(Paragraph("<b>You:</b>", bold_style))
            elements.append(Paragraph(message['content'], normal_style))
        
        elements.append(Spacer(1, 0.2 * inch))
        elements.append(Paragraph("_" * 50, normal_style))
        elements.append(Spacer(1, 0.2 * inch))

    pdf.build(elements)
    return formatted_filename

# Function to extract text from Word document
def extract_text_from_word(docx_file):
    doc = DocxDocument(docx_file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Initialize session state variables if not already present
if 'total_questions' not in st.session_state:
    st.session_state.total_questions = 3  # Control how many questions before "more questions"
if 'questions_asked' not in st.session_state:
    st.session_state.questions_asked = 0  # Tracks how many questions have been asked
if 'used_chunks' not in st.session_state:
    st.session_state.used_chunks = []  # Keep track of chunks already used for generating questions
if 'conversation' not in st.session_state:
    st.session_state.conversation = []  # Store conversation history
if 'show_next_question_button' not in st.session_state:
    st.session_state.show_next_question_button = False  # Control when to show the next question button
if 'show_input_area' not in st.session_state:
    st.session_state.show_input_area = True  # Control when to show the input area
if 'ask_more_questions' not in st.session_state:
    st.session_state.ask_more_questions = False  # Flag to ask if user wants more questions
if 'end_session' not in st.session_state:
    st.session_state.end_session = False  # Flag to check if session ended
if 'project_area' not in st.session_state:
    st.session_state.project_area = ""  # Initialize project area

# Sidebar for project area input
with st.sidebar:
    st.title("Audit Readiness")
    st.session_state.project_area = st.text_input("Enter your project area:", value=st.session_state.project_area)

    if st.button("Submit"):
        if st.session_state.project_area:
            project_area_embedding = embedding_model.embed_query(st.session_state.project_area)
            question = retrieve_chunk_and_generate_question(project_area_embedding)
            st.session_state.questions_asked += 1
            st.session_state.conversation.append({"role": "bot", "content": question})
            st.session_state.show_input_area = True
            st.rerun()  # Refresh to show the question

# Main content for conversation
st.title("🏥 Regulaider")

# Display conversation in the main area
for message in st.session_state.conversation:
    if message['role'] == 'bot':
        st.markdown(f"**Regulaider:** {message['content']}")
    else:
        st.markdown(f"**You:** {message['content']}")
    st.markdown("<hr>", unsafe_allow_html=True)

if st.session_state.conversation and st.session_state.show_input_area and not st.session_state.ask_more_questions and not st.session_state.end_session:
    # Add option for uploading a file (PDF/Word) or entering text
    input_method = st.radio("How would you like to answer?", ["Text Entry", "Upload Document"])

    if input_method == "Text Entry":
        user_response = st.text_area("Type your response here:", key="user_input")
    elif input_method == "Upload Document":
        uploaded_file = st.file_uploader("Upload File (PDF or Word)", type=["pdf", "docx"])

        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                # Extract text from PDF
                with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf:
                    user_response = ""
                    for page in pdf:
                        user_response += page.get_text()

            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Extract text from Word document
                user_response = extract_text_from_word(uploaded_file)

            st.write("Text extracted from file:")
            st.write(user_response)

    if st.button("Send"):
        if user_response:
            st.session_state.conversation.append({"role": "user", "content": user_response})
            feedback = evaluate_response(user_response, st.session_state.relevant_chunk['page_content'])
            st.session_state.conversation.append({"role": "bot", "content": feedback})

            if st.session_state.questions_asked < st.session_state.total_questions:
                st.session_state.show_input_area = False
                st.session_state.show_next_question_button = True
            else:
                st.session_state.ask_more_questions = True  # Prompt for more questions

            st.rerun()

if st.session_state.show_next_question_button:
    if st.button("Next Question"):
        project_area_embedding = embedding_model.embed_query(st.session_state.project_area)
        next_question = retrieve_chunk_and_generate_question(project_area_embedding)
        st.session_state.conversation.append({"role": "bot", "content": next_question})
        st.session_state.questions_asked += 1
        st.session_state.show_input_area = True  # Show input area again for the next question
        st.session_state.show_next_question_button = False
        st.rerun()

if st.session_state.ask_more_questions:
    st.session_state.show_input_area = False
    st.write("Do you want more questions?")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Yes!"):
            project_area_embedding = embedding_model.embed_query(st.session_state.project_area)
            next_question = retrieve_chunk_and_generate_question(project_area_embedding)
            st.session_state.conversation.append({"role": "bot", "content": next_question})
            st.session_state.show_input_area = True
            st.session_state.ask_more_questions = False
            st.rerun()

    with col2:
        if st.button("End Audit! Evaluate Responses!"):
            final_eval = generate_final_evaluation(st.session_state.conversation)
            st.session_state.conversation.append({"role": "bot", "content": final_eval})
            st.session_state.ask_more_questions = False
            st.session_state.end_session = True  # Mark session as ended
            st.rerun()

if st.session_state.end_session:
    st.session_state.show_input_area = False
    st.write("Thank you for completing the audit.")
    
    if st.button("Start New Session"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

st.sidebar.title("Download Report")

if st.sidebar.button("Download as Word"):
    filename = save_conversation_to_word(st.session_state.conversation, st.session_state.project_area)
    with open(filename, "rb") as file:
        st.sidebar.download_button(
            label="Download Word Document",
            data=file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if st.sidebar.button("Download as PDF"):
    filename = save_conversation_to_pdf(st.session_state.conversation, st.session_state.project_area)
    with open(filename, "rb") as file:
        st.sidebar.download_button(
            label="Download PDF Document",
            data=file,
            file_name=filename,
            mime="application/pdf"
        )

st.sidebar.markdown(
    """
    <style>
    .reference-text {
        position: fixed;
        bottom: 10px;
        left: 10px;
        font-size: 12px;
        color: #888888;
    }
    </style>
    <div class="reference-text">
        CFR Title 21 updated annually on April 1, 
        data available as of April 1, 2024.
    </div>
    """,
    unsafe_allow_html=True
)

# Function to save the current session to a JSON file with versioning
def save_session_to_json(project_scope, conversation):
    try:
        with open("chat_history.json", "r") as file:
            all_sessions = json.load(file)
    except FileNotFoundError:
        all_sessions = {}

    # Determine the version number for the current project scope
    version = len(all_sessions.get(project_scope, [])) + 1
    session_data = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "version": version,
        "conversation": conversation
    }
    all_sessions.setdefault(project_scope, []).append(session_data)

    # Save updated sessions back to the file
    with open("chat_history.json", "w") as file:
        json.dump(all_sessions, file, indent=4)

# Function to load saved sessions from JSON
def load_saved_sessions():
    try:
        with open("chat_history.json", "r") as file:
            return json.load(file)
    except FileNotFoundError:
        return {}

# Load previous sessions at the start
saved_sessions = load_saved_sessions()

# Display saved sessions in the sidebar with project scope and version
st.sidebar.title("Previous Sessions")
for project_scope, sessions in saved_sessions.items():
    with st.sidebar.expander(project_scope):
        for session in sessions:
            session_label = f"{project_scope} v{session['version']} - {session['timestamp']}"
            if st.sidebar.button(session_label, key=f"{project_scope}_{session['version']}"):
                # Load the selected session's conversation into the main chat area
                st.session_state.conversation = session["conversation"]
                st.session_state.project_area = project_scope
                st.session_state.end_session = False  
                st.session_state.show_input_area = False  
                st.session_state.viewing_previous_session = True  
                st.rerun()

# Save the session to JSON when ending the audit
if st.session_state.end_session:
    save_session_to_json(st.session_state.project_area, st.session_state.conversation)

if st.sidebar.button("Go to Home Page"):
    # Clear the relevant session state variables to reset to the main page
    for key in ["conversation", "project_area", "end_session", "questions_asked", "used_chunks", "show_next_question_button", "show_input_area", "ask_more_questions"]:
        st.session_state.pop(key, None)
    st.rerun()
